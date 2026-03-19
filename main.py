import os
import sys

# --- FIX: Force Gradio to use a local temp directory ---
os.environ['GRADIO_TEMP_DIR'] = os.path.join(os.getcwd(), 'gradio_temp')
if not os.path.exists(os.environ['GRADIO_TEMP_DIR']):
    os.makedirs(os.environ['GRADIO_TEMP_DIR'])
# -------------------------------------------------------

import torch
from transformers import AutoModelForSeq2SeqLM, AutoTokenizer
from IndicTransToolkit import IndicProcessor
from docx import Document
import pdfplumber
import gradio as gr

# -----------------------------
# Device setup
# -----------------------------
DEVICE = "cuda" if torch.cuda.is_available() else "cpu"

AVAILABLE_GPUS = list(range(torch.cuda.device_count()))
if not AVAILABLE_GPUS:
    raise RuntimeError("No GPUs detected. This script requires a GPU to run.")

print(f"Setup complete. Detected {len(AVAILABLE_GPUS)} GPU(s): {AVAILABLE_GPUS}")

# -----------------------------
# Language & Model Config
# -----------------------------
LANGUAGES = {
    "English": "eng_Latn",
    "Hindi": "hin_Deva",
    "Bengali": "ben_Beng",
    "Tamil": "tam_Taml",
    "Telugu": "tel_Telu",
    "Marathi": "mar_Deva",
    "Gujarati": "guj_Gujr",
    "Kannada": "kan_Knda",
    "Malayalam": "mal_Mlym",
    "Punjabi": "pan_Guru",
    "Urdu": "urd_Arab",
    "Odia": "ory_Orya",
    "Assamese": "asm_Beng",
    "Sanskrit": "san_Deva",
    "Kashmiri": "kas_Arab",
    "Sindhi": "snd_Arab",
    "Manipuri": "mni_Mtei",
    "Santali": "sat_Olch",
    "Nepali": "npi_Deva",
    "Konkani": "gom_Deva",
    "Dogri": "doi_Deva",
    "Bodo": "brx_Deva",
    "Maithili": "mai_Deva"
}

MODEL_EN_INDIC = "ai4bharat/indictrans2-en-indic-1B"
MODEL_INDIC_EN = "ai4bharat/indictrans2-indic-en-1B"
MODEL_INDIC_INDIC = "ai4bharat/indictrans2-indic-indic-1B"

loaded_models = {}

# -----------------------------
# Model Loader
# -----------------------------
def get_translation_model(src_lang_name, tgt_lang_name, gpu_id):
    global loaded_models
    
    try:
        gpu_id = int(gpu_id)
    except ValueError:
        return None, None, None, "Error: Invalid GPU ID."

    if gpu_id not in AVAILABLE_GPUS:
        return None, None, None, f"Error: GPU {gpu_id} is not available."

    if src_lang_name == "English" and tgt_lang_name != "English":
        model_key = "en-indic"
        model_name = MODEL_EN_INDIC
    elif tgt_lang_name == "English" and src_lang_name != "English":
        model_key = "indic-en"
        model_name = MODEL_INDIC_EN
    elif src_lang_name != "English" and tgt_lang_name != "English":
        model_key = "indic-indic"
        model_name = MODEL_INDIC_INDIC
    else:
        return None, None, None, "Error: Source and Target languages cannot be the same."

    cache_key = (gpu_id, model_key)
    
    if cache_key in loaded_models:
        return loaded_models[cache_key]['model'], loaded_models[cache_key]['tokenizer'], loaded_models[cache_key]['ip'], f"Model ready on GPU {gpu_id}."

    print(f"Loading model: {model_name} on GPU {gpu_id}...")
    try:
        device = f"cuda:{gpu_id}"
        
        ip = IndicProcessor(inference=True)
        tokenizer = AutoTokenizer.from_pretrained(model_name, trust_remote_code=True)
        model = AutoModelForSeq2SeqLM.from_pretrained(
            model_name, 
            trust_remote_code=True, 
            torch_dtype=torch.float16
        ).to(device)
        
        loaded_models[cache_key] = {'model': model, 'tokenizer': tokenizer, 'ip': ip}
        print(f"Model {model_key} loaded successfully.")
        return model, tokenizer, ip, f"Model loaded."
        
    except Exception as e:
        import traceback
        traceback.print_exc()
        return None, None, None, f"Error loading model: {str(e)}"

# -----------------------------
# Translation Logic (Memory Optimized & Warning Fixed)
# -----------------------------
def translate_batch_memory_safe(sentences, model, tokenizer, ip, src_lang, tgt_lang, batch_size=8):
    """
    Translates a list of sentences in small batches to avoid OOM errors.
    """
    if not sentences:
        return []

    all_translations = []
    total_sentences = len(sentences)
    
    # Process in small batches
    for i in range(0, total_sentences, batch_size):
        batch = sentences[i : i + batch_size]
        
        # Filter empty strings for processing, but keep track of indices
        valid_indices = [idx for idx, s in enumerate(batch) if s.strip()]
        valid_sentences = [batch[idx] for idx in valid_indices]
        
        if valid_sentences:
            # Preprocess
            preprocessed = ip.preprocess_batch(valid_sentences, src_lang=src_lang, tgt_lang=tgt_lang)
            
            # Tokenize
            inputs = tokenizer(
                preprocessed, 
                truncation=True, 
                padding="longest", 
                return_tensors="pt"
            ).to(model.device)

            # Generate
            with torch.no_grad():
                generated_tokens = model.generate(
                    **inputs,
                    use_cache=True,
                    min_length=0,
                    max_length=512,
                    num_beams=4,
                    early_stopping=True
                )

            # Decode (FIXED: Removed deprecated as_target_tokenizer)
            translations = tokenizer.batch_decode(
                generated_tokens.detach().cpu().tolist(), 
                skip_special_tokens=True
            )

            # Postprocess
            translations = ip.postprocess_batch(translations, lang=tgt_lang)
            
            # Place back into batch positions
            for idx, trans in zip(valid_indices, translations):
                batch[idx] = trans
        
        all_translations.extend(batch)
        
        # Clear cache to prevent memory buildup
        torch.cuda.empty_cache()
        
        if (i + batch_size) < total_sentences:
            print(f"Processed {min(i + batch_size, total_sentences)}/{total_sentences} segments...")

    return all_translations

# -----------------------------
# Structure-Preserving DOCX Processor
# -----------------------------
def process_docx(file_path, model, tokenizer, ip, src_lang, tgt_lang):
    doc = Document(file_path)
    
    # 1. Translate Paragraphs
    print("Translating paragraphs...")
    paras_text = [p.text for p in doc.paragraphs]
    # Use memory-safe translation
    translated_paras = translate_batch_memory_safe(paras_text, model, tokenizer, ip, src_lang, tgt_lang, batch_size=8)
    
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip():
            # Clear existing runs to preserve formatting style
            for run in p.runs:
                run.text = ""
            # Add translated text
            if p.runs:
                p.runs[0].text = translated_paras[i]
            else:
                p.add_run(translated_paras[i])

    # 2. Translate Tables (Cell by Cell)
    print("Translating tables...")
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    # Translate single cell content
                    translated_cell = translate_batch_memory_safe(
                        [cell.text], model, tokenizer, ip, src_lang, tgt_lang, batch_size=1
                    )[0]
                    
                    # Preserve cell formatting
                    cell.text = ""
                    for paragraph in cell.paragraphs:
                        if paragraph.text == "":
                            paragraph.add_run(translated_cell)
    
    return doc

# -----------------------------
# Main Process
# -----------------------------
def translate_document(file, src_lang_name, tgt_lang_name, gpu_id):
    if not file:
        return None, "Please upload a file."
    
    src_lang_code = LANGUAGES.get(src_lang_name)
    tgt_lang_code = LANGUAGES.get(tgt_lang_name)

    if not src_lang_code or not tgt_lang_code:
        return None, "Invalid language selection."

    model, tokenizer, ip, status = get_translation_model(src_lang_name, tgt_lang_name, gpu_id)
    if not model:
        return None, status

    input_file = file.name
    ext = os.path.splitext(input_file)[1].lower()
    
    base_name = os.path.splitext(os.path.basename(input_file))[0]
    output_file_path = os.path.join(os.getcwd(), f"{base_name}_translated_{tgt_lang_name}.docx")

    try:
        if ext == ".docx":
            print(f"Processing DOCX: {input_file}")
            doc = process_docx(input_file, model, tokenizer, ip, src_lang_code, tgt_lang_code)
            doc.save(output_file_path)
            
        elif ext == ".pdf":
            print("Processing PDF (Text Mode)...")
            text_list = []
            with pdfplumber.open(input_file) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_list.append(page_text)
            
            full_text = "\n".join(text_list)
            sentences = full_text.split('\n')
            
            # Use memory-safe translation
            translated = translate_batch_memory_safe(sentences, model, tokenizer, ip, src_lang_code, tgt_lang_code)
            
            doc = Document()
            for line in translated:
                doc.add_paragraph(line)
            doc.save(output_file_path)
            
        elif ext == ".txt":
            print("Processing TXT...")
            with open(input_file, "r", encoding="utf-8") as f:
                sentences = [line.strip() for line in f.readlines() if line.strip()]
            
            translated = translate_batch_memory_safe(sentences, model, tokenizer, ip, src_lang_code, tgt_lang_code)
            
            doc = Document()
            for line in translated:
                doc.add_paragraph(line)
            doc.save(output_file_path)
        else:
            return None, "Unsupported file type! Use DOCX, PDF, or TXT."

        return output_file_path, f"Success! File saved preserving structure."

    except Exception as e:
        import traceback
        traceback.print_exc()
        return None, f"Error: {str(e)}"

# -----------------------------
# Gradio Interface
# -----------------------------
gpu_choices = [str(i) for i in AVAILABLE_GPUS]
default_gpu = gpu_choices[0]

with gr.Blocks() as demo:
    gr.Markdown("<h1 style='text-align: center;'>Document Translation Pipeline</h1>")
    gr.Markdown("<p style='text-align: center; color:gray;'>Preserves Images, Tables, and Formatting</p>")
    
    with gr.Row():
        with gr.Column(scale=1):
            file_input = gr.File(label="Upload Document (DOCX Recommended)")
            
            with gr.Row():
                src_lang_input = gr.Dropdown(choices=list(LANGUAGES.keys()), label="Source Language", value="English")
                tgt_lang_input = gr.Dropdown(choices=list(LANGUAGES.keys()), label="Target Language", value="Hindi")
            
            gpu_input = gr.Dropdown(choices=gpu_choices, label="Select GPU", value=default_gpu)
            
            translate_button = gr.Button("Translate Document", variant="primary")
            status_output = gr.Textbox(label="Status", interactive=False)
            
        with gr.Column(scale=1):
            output_file = gr.File(label="Download Translated File")

    translate_button.click(
        fn=translate_document,
        inputs=[file_input, src_lang_input, tgt_lang_input, gpu_input],
        outputs=[output_file, status_output]
    )

if __name__ == "__main__":
    demo.launch(theme=gr.themes.Soft())
